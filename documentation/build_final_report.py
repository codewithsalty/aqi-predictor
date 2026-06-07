from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image
except Exception:
    Image = None


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "documentation"
SCREEN_DIR = DOC_DIR / "screenshots"
API_DIR = DOC_DIR / "evidence" / "api_snapshots"

DOCX_OUT = DOC_DIR / "Pearls_AQI_Predictor_Final_Internship_Report.docx"
MD_OUT = DOC_DIR / "final_submission_report.md"
INVENTORY_OUT = DOC_DIR / "screenshot_inventory.md"


NAVY = RGBColor(8, 25, 43)
TEAL = RGBColor(15, 118, 110)
SLATE = RGBColor(82, 99, 122)
LIGHT_TEAL = "DDF7F2"
LIGHT_BLUE = "EAF4F8"
LIGHT_WARN = "FFF4E5"
WHITE = "FFFFFF"


def load_json(name: str) -> dict:
    with (API_DIR / name).open("r", encoding="utf-8-sig") as f:
        return json.load(f)


prediction = load_json("live_3_day_prediction.json")
metrics = load_json("latest_model_metrics.json")
registry = load_json("latest_model_registry.json")
quality = load_json("latest_data_quality_audit.json")
pipeline = load_json("pipeline_health.json")


SCREENSHOTS = {
    "landing": SCREEN_DIR / "01_frontend_home_desktop.png",
    "dashboard": SCREEN_DIR / "02_frontend_dashboard_desktop.png",
    "methodology": SCREEN_DIR / "04_frontend_methodology.png",
    "about": SCREEN_DIR / "05_frontend_about_salman.png",
    "api_docs": SCREEN_DIR / "06_backend_fastapi_docs.png",
    "api_evidence": SCREEN_DIR / "07_backend_api_evidence.png",
    "local_artifacts": SCREEN_DIR / "08_local_artifact_evidence.png",
    "repo": SCREEN_DIR / "09_github_repo_home.png",
    "actions": SCREEN_DIR / "10_github_actions_all_green.png",
    "feature_runs": SCREEN_DIR / "11_github_feature_pipeline_runs.png",
    "training_runs": SCREEN_DIR / "12_github_training_pipeline_runs.png",
    "feature_yml": SCREEN_DIR / "13_feature_pipeline_yml_code.png",
    "training_yml": SCREEN_DIR / "14_training_pipeline_yml_code.png",
    "eda": SCREEN_DIR / "15_eda_daily_aqi_trend.png",
    "api_code": SCREEN_DIR / "16_backend_api_code.png",
    "train_code": SCREEN_DIR / "17_backend_train_code.png",
    "frontend_code": SCREEN_DIR / "18_frontend_dashboard_code.png",
    "readme": SCREEN_DIR / "19_readme_submission_proof.png",
    "atlas_overview": SCREEN_DIR / "complete screenshot of cluster 0 predict .png",
    "atlas_features": SCREEN_DIR / "features v1 (1) .png",
    "atlas_metrics": SCREEN_DIR / "model metrics.png",
    "atlas_registry": SCREEN_DIR / "model registry.png",
    "atlas_pipeline": SCREEN_DIR / "pipeline runs.png",
    "atlas_predictions": SCREEN_DIR / "predictions .png",
    "atlas_quality": SCREEN_DIR / "quality audits.png",
    "render_main": SCREEN_DIR / "render main page.png",
    "render_events": SCREEN_DIR / "render events page .png",
    "vercel": SCREEN_DIR / "vercel dashboard deployments.png",
    "vercel_logs": SCREEN_DIR / "vercel logs.png",
    "portal": SCREEN_DIR / "10 pearls project submission place github repo .png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = NAVY

    for name, size in [("Title", 30), ("Heading 1", 18), ("Heading 2", 13), ("Heading 3", 11)]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = NAVY

    styles["Heading 1"].font.color.rgb = TEAL
    styles["Heading 2"].font.color.rgb = NAVY


def add_para(doc: Document, text: str = "", bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = TEAL
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = TEAL
    r.font.size = Pt(11)
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_shading(hdr[idx], "0F766E")
        set_cell_text(hdr[idx], h, bold=True, color=RGBColor(255, 255, 255))
        if widths:
            hdr[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def picture_size(path: Path, max_width: float = 6.7, max_height: float = 4.9) -> tuple[float, float | None]:
    if Image is None:
        return max_width, None
    try:
        with Image.open(path) as img:
            w, h = img.size
        ratio = min(max_width / (w / 96), max_height / (h / 96), 1.0)
        return (w / 96) * ratio, None
    except Exception:
        return max_width, None


def add_figure(doc: Document, key: str, caption: str, max_width: float = 6.7, max_height: float = 4.9) -> None:
    path = SCREENSHOTS[key]
    if not path.exists():
        add_callout(doc, "Screenshot missing", f"Expected screenshot was not found: {path.name}", LIGHT_WARN)
        return
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(caption)
    r.bold = True
    r.font.color.rgb = TEAL
    width, _ = picture_size(path, max_width, max_height)
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def fmt(value: float | int | str, digits: int = 2) -> str:
    if isinstance(value, (int, float)):
        return f"{value:.{digits}f}"
    return str(value)


def build_docx() -> None:
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Pearls AQI Predictor")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Final Internship Project Report - Islamabad AQI Forecasting").bold = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared by Salman Khan | Data Sciences Internship | Submitted: 2026-06-08")
    add_callout(
        doc,
        "Project in one line",
        "I built a cloud-backed AQI forecasting system that collects Islamabad air quality data, engineers features, trains multiple models, stores champions in a model registry, and serves 3-day predictions through a live API and product-style web dashboard.",
        LIGHT_TEAL,
    )
    add_table(
        doc,
        ["Area", "Final implementation"],
        [
            ["City", "Islamabad, Pakistan"],
            ["Repository", "https://github.com/codewithsalty/aqi-predictor"],
            ["Frontend", "https://pearls-aqi.vercel.app/"],
            ["Dashboard", "https://pearls-aqi.vercel.app/dashboard"],
            ["Backend API", "https://aqi-predictor-api-cuec.onrender.com"],
            ["Cloud store", "MongoDB Atlas feature store, registry, predictions, quality and pipeline evidence"],
            ["Automation", "GitHub Actions scheduled and manual workflows"],
        ],
        [1.6, 5.2],
    )
    doc.add_page_break()

    doc.add_heading("1. Executive Summary", level=1)
    add_para(doc, "The internship task was to predict the next 3 days of AQI for my city using an end-to-end machine learning pipeline. I selected Islamabad because it is the city I wanted to focus on, and keeping one city made the system more reliable and easier to explain.")
    add_para(doc, "Instead of only building a notebook, I built a working product pipeline. The system collects weather and pollutant data, converts it into machine learning features, stores everything in MongoDB Atlas, trains multiple models, registers the selected champions, and exposes predictions through a FastAPI backend and a Next.js frontend.")
    add_para(doc, "The final project is not pretending that the model is perfect. It shows honest metrics, quality checks, limitations, and operational evidence. That is important because a real data science system should be explainable and maintainable, not just visually impressive.")
    add_figure(doc, "landing", "Figure 1: Live product landing page.")
    add_figure(doc, "dashboard", "Figure 2: Live dashboard with forecast, model evidence, and pipeline health.")

    doc.add_heading("2. Requirement Mapping", level=1)
    rows = [
        ["AQI next 3 days", "Completed", "Backend /predict returns Day +1, Day +2, Day +3 forecasts and risk levels."],
        ["External data API", "Completed", "Open-Meteo weather and air quality APIs are used for Islamabad."],
        ["Feature pipeline", "Completed", "GitHub Actions runs feature ingestion, feature engineering, deduplication, and Atlas writes."],
        ["Historical backfill", "Completed", "Feature store contains historical hourly Islamabad records for training."],
        ["Feature store", "Completed", "MongoDB Atlas is used as the cloud feature store."],
        ["Training pipeline", "Completed", "Training reads features from Atlas, builds targets, trains models, evaluates them, and writes registry records."],
        ["Multiple models", "Completed", "Ridge, Random Forest, Gradient Boosting, and MLP neural-net challenger are trained."],
        ["Metrics", "Completed", "RMSE, MAE, and R2 are stored in model_metrics and returned by API."],
        ["Model registry", "Completed", "Champion model metadata is stored in Atlas, with binaries in GridFS."],
        ["Automation", "Completed", "Feature pipeline is scheduled hourly; training has daily and catch-up automation."],
        ["Web dashboard", "Completed", "Next.js frontend plus FastAPI backend, deployed through Vercel and Render."],
        ["EDA and explainability", "Completed", "EDA artifacts and feature importance evidence are generated and documented."],
        ["Hazard alerts", "Completed", "Forecast outputs classify AQI risk levels, including sensitive-group warnings."],
    ]
    add_table(doc, ["Requirement", "Status", "Evidence"], rows, [1.7, 1.0, 4.1])

    doc.add_heading("3. System Architecture", level=1)
    add_para(doc, "The architecture is simple enough to defend in a viva but complete enough to show real engineering. GitHub Actions runs the pipelines, MongoDB Atlas acts as the cloud data and model store, Render hosts the backend API, and Vercel hosts the frontend.")
    add_table(
        doc,
        ["Layer", "What it does"],
        [
            ["Data source", "Fetches Islamabad weather and pollutant signals from Open-Meteo."],
            ["Feature pipeline", "Builds clean hourly records, lag features, rolling features, and time features."],
            ["Feature store", "Stores processed records in Atlas collection features_v1."],
            ["Training pipeline", "Creates 1-day, 2-day, and 3-day targets, trains all models, and evaluates them."],
            ["Model registry", "Stores champion selection, model metadata, feature list, metrics, and GridFS model artifact id."],
            ["Prediction API", "Loads latest registry and prediction records and serves JSON responses."],
            ["Frontend", "Shows forecast, model comparison, quality checks, automation health, methodology, and creator page."],
        ],
        [1.8, 5.0],
    )
    add_figure(doc, "methodology", "Figure 3: Methodology page explaining the pipeline to evaluators.")

    doc.add_heading("4. Cloud Feature Store in MongoDB Atlas", level=1)
    add_para(doc, "I used MongoDB Atlas as my cloud feature store because it was stable, visible, and easy to prove during evaluation. All feature records are stored in the aqi_predictor database, mainly inside features_v1.")
    add_para(doc, "Important point: I did not store secrets in the code. The MongoDB URI is kept in GitHub Actions secrets and Render environment variables.")
    add_table(
        doc,
        ["Collection", "Purpose"],
        [
            ["features_v1", "Processed Islamabad AQI and pollutant feature rows."],
            ["model_metrics", "Evaluation metrics for all trained models and forecast horizons."],
            ["model_registry", "Champion models, feature list, GridFS id, and leaderboard summary."],
            ["predictions", "Generated 3-day forecasts and risk classes."],
            ["pipeline_runs", "Success/failure logs from feature, training, and audit runs."],
            ["quality_audits", "Null, duplicate, AQI range, and leakage checks."],
            ["fs.files / fs.chunks", "GridFS model artifact storage."],
        ],
        [1.8, 5.0],
    )
    add_figure(doc, "atlas_overview", "Figure 4: MongoDB Atlas database collections used by the project.")
    add_figure(doc, "atlas_features", "Figure 5: features_v1 collection with Islamabad cloud feature records.")

    doc.add_heading("5. Feature Engineering", level=1)
    add_para(doc, "Raw data is useful, but models need structured signals. The feature pipeline converts the raw AQI and pollutant readings into time-aware features so the model can learn movement and recent history.")
    add_table(
        doc,
        ["Feature group", "Examples"],
        [
            ["Pollutants", "pm10, pm2_5, co, no2, so2, o3"],
            ["Time features", "day, month"],
            ["AQI movement", "aqi_delta"],
            ["Lag features", "aqi_lag_1, aqi_lag_2, aqi_lag_3, pm2_5_lag_1, pm10_lag_1"],
            ["Rolling features", "pm2_5_roll_3, pm10_roll_3"],
        ],
        [1.8, 5.0],
    )
    add_para(doc, "The stored feature list in the latest registry contains 17 columns, including pollutant values, time features, lag features, rolling features, and AQI delta.")

    doc.add_heading("6. Training Pipeline and Model Selection", level=1)
    add_para(doc, "The training pipeline reads historical features from Atlas. It creates future targets for Day +1, Day +2, and Day +3, then trains all candidate models for each forecast horizon.")
    add_para(doc, "The champion is not hardcoded. For each horizon, the model with the best validation behavior is selected and stored. The system also stores an overall leaderboard for model comparison.")
    add_table(
        doc,
        ["Model", "Why included"],
        [
            ["Ridge Regression", "Strong baseline, fast, less likely to overfit."],
            ["Random Forest", "Handles non-linear pollutant and weather relationships well."],
            ["Gradient Boosting", "Competitive tree-based challenger."],
            ["MLP Neural Net", "Neural-network style challenger using tabular engineered features."],
        ],
        [1.8, 5.0],
    )
    champion_rows = []
    for day_key, champion in registry["champion_model"].items():
        best = next(row for row in metrics["metrics"][day_key] if row["model"] == champion)
        champion_rows.append([day_key.replace("_", " ").title(), champion, fmt(best["rmse"]), fmt(best["mae"]), fmt(best["r2"], 3)])
    add_table(doc, ["Horizon", "Champion", "RMSE", "MAE", "R2"], champion_rows, [1.2, 1.7, 1.1, 1.1, 1.1])
    add_para(doc, f"The overall leaderboard winner is {registry['overall_winner']['model']}, with average RMSE {fmt(registry['overall_winner']['avg_rmse'])}, average MAE {fmt(registry['overall_winner']['avg_mae'])}, and average R2 {fmt(registry['overall_winner']['avg_r2'], 3)}.")
    add_figure(doc, "atlas_metrics", "Figure 6: Model metrics stored in MongoDB Atlas.")
    add_figure(doc, "atlas_registry", "Figure 7: Model registry entries stored in Atlas with champion metadata and GridFS artifact reference.")

    doc.add_heading("7. Prediction Output", level=1)
    add_para(doc, "The backend returns the latest 3-day forecast. The default mode uses horizon champions, meaning each day can use the best model selected for that specific forecast horizon. This is why Day +1 can use Ridge while Day +2 and Day +3 can use Random Forest.")
    pred_rows = []
    for idx, pred in enumerate(prediction["predictions"], start=1):
        model = prediction["model"][f"day_{idx}"]
        pred_rows.append([f"Day +{idx}", pred["date"], fmt(pred["aqi"]), pred["risk"], model])
    add_table(doc, ["Horizon", "Date", "Predicted AQI", "Risk", "Model"], pred_rows, [1.0, 1.2, 1.3, 2.0, 1.3])
    add_para(doc, "The API also supports model override, so the dashboard can compare predictions from a selected model instead of only showing the automatic champion.")
    add_figure(doc, "atlas_predictions", "Figure 8: Prediction records stored in MongoDB Atlas.")
    add_figure(doc, "api_evidence", "Figure 9: Live backend API evidence page.")

    doc.add_heading("8. Automation with GitHub Actions", level=1)
    add_para(doc, "Automation was one of the biggest parts of this project. I configured GitHub Actions so the pipelines can run without me manually opening the terminal.")
    add_table(
        doc,
        ["Workflow", "Schedule", "What happens"],
        [
            ["Feature Pipeline", "Hourly with backup cron", "Fetches recent Islamabad data, engineers features, deduplicates rows, stores cloud feature records, and logs the run."],
            ["Training Pipeline", "Daily with backup cron", "Trains models when due, evaluates them, updates registry, stores artifacts, and writes predictions."],
            ["Manual Recovery", "Manual trigger", "Allows safe reruns if GitHub scheduled runners are delayed."],
        ],
        [1.5, 1.6, 3.7],
    )
    add_callout(
        doc,
        "Why there are backup triggers",
        "GitHub cron can be delayed on free/public runners, so I added backup schedule triggers and a training catch-up check. This makes the automation more robust while deduplication prevents duplicate feature rows from corrupting the database.",
        LIGHT_TEAL,
    )
    add_figure(doc, "actions", "Figure 10: GitHub Actions evidence showing green automated workflow runs.")
    add_figure(doc, "feature_runs", "Figure 11: Feature pipeline scheduled/manual run history.")
    add_figure(doc, "training_runs", "Figure 12: Training pipeline scheduled/manual run history.")
    add_figure(doc, "feature_yml", "Figure 13: Feature pipeline workflow code.")
    add_figure(doc, "training_yml", "Figure 14: Training pipeline workflow code.")

    doc.add_heading("9. Data Quality, EDA, and Explainability", level=1)
    add_para(doc, "I added quality checks because AQI forecasting can fail silently if the data has duplicates, nulls, or leakage. The latest quality audit status is pass.")
    add_table(
        doc,
        ["Check", "Latest value"],
        [
            ["Rows audited", str(quality["row_count"])],
            ["Duplicate rows", str(quality["duplicate_rows"])],
            ["AQI out of range", str(quality["aqi_out_of_range_rows"])],
            ["Leakage risk hint", str(quality["leakage_risk_hint"])],
            ["Status", quality["status"]],
        ],
        [2.2, 4.6],
    )
    add_para(doc, "The project also produces EDA artifacts such as summary statistics, correlations, daily AQI trend visualizations, and feature-importance style evidence. These artifacts help explain what the model is learning and make the project stronger than a black-box dashboard.")
    add_figure(doc, "atlas_quality", "Figure 15: Quality audit collection with pass status.")
    add_figure(doc, "eda", "Figure 16: EDA daily AQI trend artifact.")
    add_figure(doc, "local_artifacts", "Figure 17: Local EDA, SHAP/importance, and artifact evidence.")

    doc.add_heading("10. Backend API and Deployment", level=1)
    add_para(doc, "The backend is a FastAPI service deployed on Render. It reads from Atlas and exposes clean endpoints for the frontend and for evaluators.")
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["/health", "Health check for service and database connection."],
            ["/predict", "Latest 3-day AQI forecast and risk levels."],
            ["/metrics/latest", "Latest model metrics and leaderboard."],
            ["/models/latest", "Latest registry metadata and champion models."],
            ["/pipeline/health", "Recent pipeline runs and GitHub Actions evidence."],
            ["/quality/latest", "Latest data quality audit."],
        ],
        [1.6, 5.2],
    )
    add_figure(doc, "api_docs", "Figure 18: FastAPI documentation generated by the live backend.")
    add_figure(doc, "render_main", "Figure 19: Render backend service overview.")
    add_figure(doc, "render_events", "Figure 20: Render deployment events showing successful backend deployment.")

    doc.add_heading("11. Frontend Product", level=1)
    add_para(doc, "The frontend is not only a plain chart page. It is structured as a product website with a landing page, dashboard, methodology page, and about page. The dashboard pulls from the backend API and shows forecasts, model comparison, registry evidence, quality status, and pipeline health.")
    add_para(doc, "I also included a creator/about page so the submission has a human face and can be used as a portfolio project later.")
    add_figure(doc, "dashboard", "Figure 21: Dashboard connected to the backend prediction API.")
    add_figure(doc, "about", "Figure 22: About Salman page for portfolio/presentation use.")
    add_figure(doc, "vercel", "Figure 23: Vercel production deployment evidence.")
    add_figure(doc, "vercel_logs", "Figure 24: Vercel deployment logs evidence.")

    doc.add_heading("12. Code Organization and Security", level=1)
    add_para(doc, "The repository is organized so an evaluator can understand the work quickly. Backend, frontend, workflows, docs, evidence, and artifacts are separated.")
    add_table(
        doc,
        ["Folder/file", "Purpose"],
        [
            ["backend/app", "FastAPI app, training, feature engineering, database access, model logic."],
            ["backend/scripts", "Command entry points used by GitHub Actions."],
            ["frontend/app", "Next.js pages and dashboard UI."],
            [".github/workflows", "Feature, training, and recovery automation."],
            ["documentation", "Final report, screenshots, evidence, and API snapshots."],
            ["backend/artifacts", "EDA and local evidence artifacts."],
        ],
        [1.8, 5.0],
    )
    add_para(doc, "Security decisions I followed:")
    add_bullets(doc, [
        "No MongoDB URI or password is committed in source code.",
        "GitHub Actions uses repository secrets.",
        "Render uses environment variables.",
        "The frontend only calls public backend endpoints and does not expose database credentials.",
        "The report uses screenshots and public-safe evidence instead of secret values.",
    ])
    add_figure(doc, "repo", "Figure 25: Public GitHub repository prepared for submission.")
    add_figure(doc, "api_code", "Figure 26: Backend API implementation evidence.")
    add_figure(doc, "train_code", "Figure 27: Training and champion-selection implementation evidence.")
    add_figure(doc, "frontend_code", "Figure 28: Frontend dashboard implementation evidence.")

    doc.add_heading("13. Challenges and Fixes", level=1)
    add_table(
        doc,
        ["Challenge", "How I fixed it"],
        [
            ["GitHub account billing lock", "Created/used the clean codewithsalty repo so Actions could run normally."],
            ["MongoDB URI failed first time", "Fixed password URL encoding and kept the URI only in secrets."],
            ["Old city appeared as Lahore", "Reconfigured the project for Islamabad and reran ingestion/training."],
            ["Scheduled workflows looked delayed", "Added backup cron schedules and training catch-up logic."],
            ["Frontend initially could not fetch API", "Connected frontend environment to the Render backend and added CORS support."],
            ["Model results varied by horizon", "Stored horizon-level champions instead of forcing one model for all three days."],
        ],
        [2.2, 4.6],
    )
    add_para(doc, "These issues were useful because they made the final project more realistic. I had to solve cloud, automation, deployment, and data problems, not just model training.")

    doc.add_heading("14. Why This Project Stands Out", level=1)
    add_bullets(doc, [
        "It is an end-to-end system, not only a notebook.",
        "It has cloud evidence in Atlas, GitHub Actions, Render, and Vercel.",
        "It trains multiple models and selects champions dynamically.",
        "It stores metrics, registry metadata, predictions, and pipeline logs.",
        "It includes data quality checks and honest model limitations.",
        "It has a product-style frontend that makes the system easy to understand.",
        "It is modular enough to expand from Islamabad to more cities later.",
    ])

    doc.add_heading("15. Final Submission Evidence", level=1)
    add_para(doc, "The candidate portal asks for the GitHub repository URL. The repository is public and includes code, workflows, README, evidence, documentation, backend/frontend implementation, and deployment links.")
    add_table(
        doc,
        ["Submission item", "Link/value"],
        [
            ["GitHub repository", "https://github.com/codewithsalty/aqi-predictor"],
            ["Frontend live app", "https://pearls-aqi.vercel.app/"],
            ["Dashboard", "https://pearls-aqi.vercel.app/dashboard"],
            ["Backend API", "https://aqi-predictor-api-cuec.onrender.com"],
            ["Backend API docs", "https://aqi-predictor-api-cuec.onrender.com/docs"],
            ["Database", "MongoDB Atlas aqi_predictor database, evidence included in screenshots"],
        ],
        [1.7, 5.1],
    )
    add_figure(doc, "readme", "Figure 29: README submission proof in the repository.")
    add_figure(doc, "portal", "Figure 30: Candidate portal submission area with GitHub repository URL.")

    doc.add_heading("16. Conclusion", level=1)
    add_para(doc, "This project gave me a full view of what a real machine learning product looks like. The model is only one part. The bigger work is data collection, feature storage, automation, quality checks, API reliability, UI clarity, and proof that the system works.")
    add_para(doc, "My final submission is a complete AQI forecasting system for Islamabad with cloud-backed evidence and a live interface. It is ready for internship evaluation and also strong enough to keep improving as a portfolio project.")

    # Footer
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Pearls AQI Predictor - Salman Khan - Final Internship Report").font.size = Pt(8)

    doc.save(DOCX_OUT)


def build_markdown() -> None:
    pred_lines = []
    for idx, pred in enumerate(prediction["predictions"], start=1):
        model = prediction["model"][f"day_{idx}"]
        pred_lines.append(f"| Day +{idx} | {pred['date']} | {fmt(pred['aqi'])} | {pred['risk']} | {model} |")

    champ_lines = []
    for day_key, champion in registry["champion_model"].items():
        best = next(row for row in metrics["metrics"][day_key] if row["model"] == champion)
        champ_lines.append(f"| {day_key.replace('_', ' ').title()} | {champion} | {fmt(best['rmse'])} | {fmt(best['mae'])} | {fmt(best['r2'], 3)} |")

    text = f"""# Pearls AQI Predictor - Final Internship Report

Prepared by Salman Khan  
City: Islamabad, Pakistan  
Submission date: 2026-06-08

## Executive Summary

I built an end-to-end AQI forecasting system for Islamabad that predicts the next 3 days of AQI. The project uses a cloud feature store, automated model training, a model registry, a backend API, and a live frontend dashboard.

This was not only a notebook. The final system includes:

- Open-Meteo weather and air-quality ingestion.
- MongoDB Atlas feature store and model registry.
- Feature engineering with lag, rolling, time, pollutant, and AQI movement features.
- Ridge, Random Forest, Gradient Boosting, and MLP neural-net challenger models.
- RMSE, MAE, and R2 evaluation.
- GitHub Actions automation.
- Render backend deployment.
- Vercel frontend deployment.
- EDA, quality audit, and evidence screenshots.

## Final Links

| Item | Link |
|---|---|
| GitHub repository | https://github.com/codewithsalty/aqi-predictor |
| Frontend | https://pearls-aqi.vercel.app/ |
| Dashboard | https://pearls-aqi.vercel.app/dashboard |
| Backend API | https://aqi-predictor-api-cuec.onrender.com |
| Backend docs | https://aqi-predictor-api-cuec.onrender.com/docs |

## Latest Prediction Snapshot

Generated at: {prediction['generated_at']}  
Selection mode: {prediction['selection_mode']}

| Horizon | Date | Predicted AQI | Risk | Model |
|---|---:|---:|---|---|
{chr(10).join(pred_lines)}

## Latest Champion Models

| Horizon | Champion | RMSE | MAE | R2 |
|---|---|---:|---:|---:|
{chr(10).join(champ_lines)}

Overall winner: {registry['overall_winner']['model']}  
Average RMSE: {fmt(registry['overall_winner']['avg_rmse'])}  
Average MAE: {fmt(registry['overall_winner']['avg_mae'])}  
Average R2: {fmt(registry['overall_winner']['avg_r2'], 3)}

## Quality Audit

| Check | Value |
|---|---:|
| Rows audited | {quality['row_count']} |
| Duplicate rows | {quality['duplicate_rows']} |
| AQI out of range rows | {quality['aqi_out_of_range_rows']} |
| Leakage risk hint | {quality['leakage_risk_hint']} |
| Status | {quality['status']} |

## What I Submitted

The candidate portal requires the GitHub repository URL. The repository is public and includes the complete backend, frontend, automation workflows, documentation, screenshots, evidence files, and deployment information.

The complete polished report is available as:

`documentation/Pearls_AQI_Predictor_Final_Internship_Report.docx`
"""
    MD_OUT.write_text(text, encoding="utf-8")


def build_inventory() -> None:
    rows = [
        ("01_frontend_home_desktop.png", "Live landing page"),
        ("02_frontend_dashboard_desktop.png", "Live dashboard"),
        ("04_frontend_methodology.png", "Methodology page"),
        ("05_frontend_about_salman.png", "About Salman page"),
        ("06_backend_fastapi_docs.png", "FastAPI docs"),
        ("07_backend_api_evidence.png", "Live API evidence"),
        ("08_local_artifact_evidence.png", "EDA/model artifact evidence"),
        ("09_github_repo_home.png", "GitHub repository"),
        ("10_github_actions_all_green.png", "Green Actions evidence"),
        ("11_github_feature_pipeline_runs.png", "Feature pipeline runs"),
        ("12_github_training_pipeline_runs.png", "Training pipeline runs"),
        ("13_feature_pipeline_yml_code.png", "Feature workflow code"),
        ("14_training_pipeline_yml_code.png", "Training workflow code"),
        ("15_eda_daily_aqi_trend.png", "EDA trend plot"),
        ("16_backend_api_code.png", "Backend API code"),
        ("17_backend_train_code.png", "Training code"),
        ("18_frontend_dashboard_code.png", "Frontend dashboard code"),
        ("19_readme_submission_proof.png", "README proof"),
        ("complete screenshot of cluster 0 predict .png", "MongoDB Atlas collections overview"),
        ("features v1 (1) .png", "Atlas feature store rows"),
        ("model metrics.png", "Atlas model metrics"),
        ("model registry.png", "Atlas model registry"),
        ("pipeline runs.png", "Atlas pipeline run logs"),
        ("predictions .png", "Atlas prediction records"),
        ("quality audits.png", "Atlas quality audits"),
        ("render main page.png", "Render backend service"),
        ("render events page .png", "Render deploy events"),
        ("vercel dashboard deployments.png", "Vercel production deploy"),
        ("vercel logs.png", "Vercel deployment logs"),
        ("10 pearls project submission place github repo .png", "Candidate portal submission evidence"),
    ]
    lines = [
        "# Submission Screenshot Inventory",
        "",
        "This folder contains the final evidence screenshots used in the Pearls AQI Predictor report.",
        "",
        "The mobile screenshot was intentionally not used in the final report because the requested submission document focuses on desktop/product and cloud evidence.",
        "",
        "| Screenshot | What it proves |",
        "|---|---|",
    ]
    for name, purpose in rows:
        lines.append(f"| `{name}` | {purpose} |")
    INVENTORY_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    build_inventory()
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {MD_OUT}")
    print(f"Wrote {INVENTORY_OUT}")
